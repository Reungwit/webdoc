from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import os

# ⬇️ ใช้ยูทิลิตี้จากไฟล์กลาง
from .doc_function import *

# หน้าปกไทย
def doc_cover_th(project_name_th, project_name_en,
                 author1_th, author2_th,
                 author1_en, author2_en,
                 academic_year_be, dep_th):

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ตั้งค่าขอบกระดาษ
    section = doc.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    width_in_inches = 5 / 2.54  # 5 cm to inches
    height_in_inches = 5 / 2.54

    logo_path = os.path.join('static', 'img', 'kmutnb_logo_cover.png')

    # แทรกรูปภาพ
    doc.add_picture(logo_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # เนื้อความหน้าปก (คงไว้ในไฟล์นี้)
    doc.add_paragraph("").alignment = 1
    title = doc.add_paragraph(project_name_th)
    title.alignment = 1
    doc.add_paragraph(project_name_en).alignment = 1

    doc.add_paragraph("\n\n\n\n\n\n")
    doc.add_paragraph(f" {author1_th}").alignment = 1
    doc.add_paragraph(f" {author2_th}").alignment = 1
    doc.add_paragraph("\n\n\n")
    doc.add_paragraph("ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรอุตสาหกรรมศาสตรบัณฑิต").alignment = 1
    doc.add_paragraph("สาขาวิชาเทคโนโลยีสารสนเทศ ภาควิชาเทคโนโลยีสารสนเทศ").alignment = 1
    doc.add_paragraph(f"{dep_th}").alignment = 1
    doc.add_paragraph("มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1
    doc.add_paragraph(f"ปีการศึกษา {academic_year_be}").alignment = 1
    doc.add_paragraph("ลิขสิทธิ์ของมหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1

    return doc

# หน้าปกรอง (ไทย)
def doc_cover_sec(project_name_th, project_name_en,
                  author1_th, author2_th,
                  author1_en, author2_en,
                  academic_year, dep_th):

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    doc.add_paragraph("").alignment = 1
    title = doc.add_paragraph(project_name_th)
    title.alignment = 1
    doc.add_paragraph(project_name_en).alignment = 1

    doc.add_paragraph("\n\n\n\n\n\n\n")
    doc.add_paragraph(f" {author1_th}").alignment = 1
    doc.add_paragraph(f" {author2_th}").alignment = 1
    doc.add_paragraph("\n\n\n\n\n\n\n")
    doc.add_paragraph("ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรอุตสาหกรรมศาสตรบัณฑิต").alignment = 1
    doc.add_paragraph("สาขาวิชาเทคโนโลยีสารสนเทศ ภาควิชาเทคโนโลยีสารสนเทศ").alignment = 1
    doc.add_paragraph(f"{dep_th}").alignment = 1
    doc.add_paragraph("มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1
    doc.add_paragraph(f"ปีการศึกษา {academic_year}").alignment = 1
    doc.add_paragraph("ลิขสิทธิ์ของมหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1

    return doc

# หน้าปกภาษาอังกฤษ
def doc_cover_en(project_name_th, project_name_en,
                 author1_th, author2_th,
                 author1_en, author2_en,
                 academic_year, dep_en):

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    doc.add_paragraph("").alignment = 1
    title = doc.add_paragraph(project_name_en)
    title.alignment = 1

    doc.add_paragraph("\n\n\n\n\n\n\n")
    doc.add_paragraph(f"{author1_en}").alignment = 1
    doc.add_paragraph(f"{author2_en}").alignment = 1
    doc.add_paragraph("\n\n\n\n\n\n\n")
    doc.add_paragraph("PROJECT REPORT SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS").alignment = 1
    doc.add_paragraph("FOR THE BACHELOR’S DEGREE OF INDUSTRIAL TECHNOLOGY").alignment = 1
    doc.add_paragraph("PROGRAM IN INFORMATION TECHNOLOGY").alignment = 1
    doc.add_paragraph("DEPARTMENT OF INFORMATION TECHNOLOGY").alignment = 1
    doc.add_paragraph(f"{dep_en}").alignment = 1
    doc.add_paragraph("KING MONGKUT'S UNIVERSITY OF TECHNOLOGY NORTH BANGKOK").alignment = 1

    # แปลงปี พ.ศ. → ค.ศ. (ถ้าใส่มาเป็น พ.ศ.)
    try:
        academic_year_int = int(academic_year)
        gregorian_year = academic_year_int - 543
        doc.add_paragraph(f"ACADEMIC YEAR {gregorian_year}").alignment = 1
    except:
        doc.add_paragraph("ACADEMIC YEAR (____)").alignment = 1

    doc.add_paragraph("COPYRIGHT OF KING MONGKUT'S UNIVERSITY OF TECHNOLOGY NORTH BANGKOK").alignment = 1

    return doc

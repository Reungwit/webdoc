# -*- coding: utf-8 -*-
from __future__ import annotations
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

from man_doc.doc_function import (
    doc_setup,
    apply_rest_page_margin,
    add_center_paragraph,
    add_body_paragraph_style_1,
    add_section_heading_level1_style_1,
    add_section_heading_level2_plus_style_1,
    add_picture_box_with_caption,
    walk_item_tree,
    resolve_image_path,
    t, as_list,
    add_intro_caption_paragraph,
    # === นำวิธีจากบทที่ 2 มาใช้ ===
    make_heading_tap_func_map,
    add_wrapped_paragraph,
)

CHAPTER_NO = 3   # เลขบทสำหรับรันเลขรูป: ภาพที่ 3-1, 3-2, ...


def get_table_no(tinfo: Dict[str, Any], idx: int, chapter_no: int = CHAPTER_NO) -> str:
    """
    ดึงเลขตารางจาก JSON
    - ถ้ามี field table_no (เช่น '3-1') ก็ใช้เลย
    - ถ้าไม่มีให้สร้างจากเลขบท + ลำดับ idx -> '3-idx'
    """
    raw = str(tinfo.get("table_no") or "").strip()
    if raw:
        return raw  # ใช้เลขที่มาจาก frontend/db (computeFirstFreeTableNoForChapter)
    return f"{chapter_no}-{idx}"


def add_tables_from_json(doc: Document, tables_json, font_name="TH SarabunPSK", base_pt=16):
    """วาดตารางจากโครงสร้าง [{caption, headers, rows, table_no, ...}]"""
    if not tables_json:
        return

    for idx, tinfo in enumerate(tables_json, start=1):
        headers = tinfo.get("headers") or []
        rows    = tinfo.get("rows") or []
        caption = tinfo.get("caption") or ""

        # ✅ ดึงเลขตารางจาก JSON (เช่น 3-1) หรือสร้างใหม่ถ้าไม่มี
        table_no = get_table_no(tinfo, idx)
        ncols = len(headers) or (len(rows[0]) if rows else 0)
        ncols = max(1, ncols)

        # ตารางจริง
        table = doc.add_table(rows=len(rows)+1, cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # ส่วนหัว
        for ci in range(ncols):
            text = headers[ci] if ci < len(headers) else f"หัว {ci+1}"
            cell = table.cell(0, ci)
            cell.text = str(text)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(base_pt)

        # เนื้อหาตาราง
        for ri, row in enumerate(rows, start=1):
            for ci in range(ncols):
                v = row[ci] if ci < len(row) else ""
                cell = table.cell(ri, ci)
                cell.text = str(v)
                for run in cell.paragraphs[0].runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(base_pt)
                    
        # คำอธิบายตาราง (caption)
        full_caption_no  = f"ตารางที่ {table_no}"
        full_caption_txt = f"{caption}"

        if caption or table_no:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- ตัวหนาเฉพาะ "ตารางที่ x-x" ---
            r1 = p.add_run(full_caption_no + "  ")   # เว้น 2 ช่อง
            r1.bold = True
            r1.font.name = font_name
            r1._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            r1.font.size = Pt(base_pt)

            # --- ตัวบางเฉพาะ caption ---
            r2 = p.add_run(full_caption_txt)
            r2.bold = False
            r2.font.name = font_name
            r2._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            r2.font.size = Pt(base_pt)

       
       
        # "จากตารางที่ เลขบท-เลขตาราง  คำอธิบายตาราง"
        for raw_note in as_list(tinfo.get("notes")):
            note = t(raw_note)
            if note:
                add_wrapped_paragraph(
                    doc,
                    f"จากตารางที่ {table_no}  {note}",
                    n=99999,
                    custom_tap=0.75,
                    disth=True,
                )  # ใช้สไตล์เดียวกับคำอธิบายรูป :contentReference[oaicite:1]{index=1}


        # เว้นบรรทัดหลังตาราง
        doc.add_paragraph("")


# -------------------------- helper ภายในบทที่ 3 --------------------------

def _render_intro(
    doc: Document,
    intro_any: Any,
    *,
    media_root: str,
    chapter_no: int,
    pic_counter: List[int],
    seen_pics: set[str],
    heading_func,
) -> None:
    """
    intro_any รองรับ:
      - str  : ย่อหน้าเดียว
      - dict : {"paragraphs":[...], "items":[...], "pictures":[...]}
    ใช้สไตล์ย่อหน้า/รูป/คำอธิบายใต้รูปแบบเดียวกับบทที่ 2
    """
    # กรณีเป็น dict ตามสเปคใหม่
    if isinstance(intro_any, dict):
        # ย่อหน้า intro
        for raw in as_list(intro_any.get("paragraphs")):
            s = t(raw)
            if s:
                # ใช้ add_wrapped_paragraph แบบบทที่ 2 (ย่อหน้า intro)
                add_wrapped_paragraph(doc, s, n=99999, custom_tap=0.75, disth=True)
                # --- subnodes: หัวข้อย่อยในบทนำ เช่น 3.1, 3.2 ---
        for idx, sn in enumerate(as_list(intro_any.get("subnodes"))):
            title_txt = t((sn or {}).get("title"))
            title_no = f"{chapter_no}.{idx+1}"

            if title_no or title_txt:
                add_section_heading_level2_plus_style_1(doc, title_no, title_txt)

            for p in as_list((sn or {}).get("paragraphs")):
                s = t(p)
                if s:
                    add_wrapped_paragraph(doc, s, n=99999, custom_tap=0.75, disth=True)

        # รูปในบทนำ
        for pinfo in as_list(intro_any.get("pictures")):
            abs_path = resolve_image_path(pinfo, media_root)
            if not abs_path or abs_path in seen_pics:
                continue
            seen_pics.add(abs_path)
            pic_counter[0] += 1
            run_no = pic_counter[0]

            add_picture_box_with_caption(
                doc,
                abs_path,
                pic_name=t((pinfo or {}).get("pic_name")),
                chapter_no=chapter_no,
                run_no=run_no,
            )
            # คำอธิบายใต้รูปแบบบทที่ 2: "จากภาพที่ 3-x  ...."
            for cap in as_list(pinfo.get("captions")):
                s = t(cap)
                if s:
                    add_intro_caption_paragraph(
                        doc,
                        f"จากภาพที่ {chapter_no}-{run_no}  {s}"
                    )

        # items ใต้บทนำ (ใช้ walk_item_tree + heading_func แบบบทที่ 2)
        items = as_list(intro_any.get("items"))
        if items:
            walk_item_tree(
                doc,
                section_no="",
                nodes=items,
                chapter_no=chapter_no,
                media_root=media_root,
                pic_counter=pic_counter,
                seen_pics=seen_pics,
                heading_func=heading_func,
                body_func=add_body_paragraph_style_1,
                caption_func=lambda d, s: add_intro_caption_paragraph(
                    d,
                    f"จากภาพที่ {chapter_no}-{pic_counter[0]}  {t(s)}",
                ),
            )
        return

    # กรณี intro เป็น string ธรรมดา
    s = t(intro_any)
    if s:
        add_wrapped_paragraph(doc, s, n=99999, custom_tap=0.75, disth=True)


def _render_sections(
    doc: Document,
    sections_json: List[Dict[str, Any]] | None,
    *,
    media_root: str,
    chapter_no: int,
    heading_func,
    pic_counter: List[int],
    seen_pics: set[str],
) -> None:
    """
    เรนเดอร์หัวข้อหลัก/หัวข้อย่อย + รูป/คำอธิบายใต้รูป
    โดยยึดวิธีเดียวกับ doc_chapter2
    """
    sections_json = sections_json or []

    for sec in sections_json:
        sec_no    = t((sec or {}).get("title_no"))
        sec_title = t((sec or {}).get("title"))

        # หัวข้อระดับ 1
        if sec_no or sec_title:
            add_section_heading_level1_style_1(doc, sec_no, sec_title)

        # เนื้อหาใต้หัวข้อ
        for raw in as_list((sec or {}).get("body_paragraphs")):
            s = t(raw)
            if s:
                add_body_paragraph_style_1(doc, s)

        # รูประดับหัวข้อ + คำอธิบายใต้รูปตามแบบบทที่ 2
        for pinfo in as_list((sec or {}).get("pictures")):
            abs_path = resolve_image_path(pinfo, media_root)
            if not abs_path or abs_path in seen_pics:
                continue
            seen_pics.add(abs_path)
            pic_counter[0] += 1
            run_no = pic_counter[0]

            add_picture_box_with_caption(
                doc,
                abs_path,
                pic_name=t((pinfo or {}).get("pic_name")),
                chapter_no=chapter_no,
                run_no=run_no,
            )
            for cap in as_list(pinfo.get("captions")):
                s = t(cap)
                if s:
                    add_intro_caption_paragraph(
                        doc,
                        f"จากภาพที่ {chapter_no}-{run_no}  {s}",
                    )

        # หัวข้อย่อยเป็นต้นไม้: ใช้ heading_func + caption_func แบบบทที่ 2
        walk_item_tree(
            doc,
            sec_no,
            as_list((sec or {}).get("items")),
            chapter_no=chapter_no,
            media_root=media_root,
            pic_counter=pic_counter,
            seen_pics=seen_pics,
            heading_func=heading_func,
            body_func=add_body_paragraph_style_1,
            caption_func=lambda d, s: add_intro_caption_paragraph(
                d,
                f"จากภาพที่ {chapter_no}-{pic_counter[0]}  {t(s)}",
            ),
        )


def doc_chapter3(
    intro_body: Any = "",
    sections_json: List[Dict[str, Any]] | None = None,
    tables_json: List[Dict[str, Any]] | None = None,
    media_root: str = "",
) -> Document:
    """
    ตัวสร้างไฟล์บทที่ 3
    - intro_body  : str หรือ dict {"paragraphs":[...], "items":[...], "pictures":[...]}
    - sections_json: โครงสร้างหัวข้อ/ย่อหน้าย่อย
    - tables_json : [{caption, headers, rows, ...}]
    """
    doc = doc_setup()

    # หัวบท
    add_center_paragraph(doc, "บทที่ 3", bold=True, font_size=20)
    add_center_paragraph(doc, "วิธีการดำเนินงาน\n", bold=True, font_size=20)

    # ปรับ margin หน้าถัดไป
    apply_rest_page_margin(doc, top_inch=1.5)

    # -------- สร้าง heading_func แบบเดียวกับบทที่ 2 --------
    heading_func = make_heading_tap_func_map(
        level_tap_cm={
            1: 0.00,  # 3.x
            2: 0.75,  # 3.x.x
            3: 1.95,  # 3.x.x.x
            4: 3.35,  # 3.x.x.x.x
            5: 5.05,  # level 5 (ก) ข) ...)
        },
        level_bold={
            1: True,
            2: False,
            3: False,
            4: False,
            5: False,
        },
        alpha_level=5,       # level 5 แสดง ก) ข) ...
        # alpha_list=THAI_ALPHA  # ถ้าต้องการ override
        fallback_left_cm=0.0,
    )

    # ตัวนับรูปใช้ร่วมกันทั้งบท (intro + sections + items)
    pic_counter = [0]
    seen_pics: set[str] = set()

    # บทนำ
    _render_intro(
        doc,
        intro_body,
        media_root=media_root,
        chapter_no=CHAPTER_NO,
        pic_counter=pic_counter,
        seen_pics=seen_pics,
        heading_func=heading_func,
    )

    # หัวข้อใหญ่/ย่อย
    _render_sections(
        doc,
        sections_json or [],
        media_root=media_root,
        chapter_no=CHAPTER_NO,
        heading_func=heading_func,
        pic_counter=pic_counter,
        seen_pics=seen_pics,
    )

    # ตาราง (ไม่เปลี่ยน logic เดิม)
    add_tables_from_json(doc, tables_json)

    return doc

# aliases
doc_chapter_3 = doc_chapter3
generate_doc = doc_chapter3

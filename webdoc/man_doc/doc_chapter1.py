from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document as DocxDocument
from pythainlp.tokenize import word_tokenize

# 🔹 เรียกใช้ฟังก์ชันจากไฟล์ doc_function.py
from .doc_function import (
    doc_setup,
    add_center_paragraph,
    add_left_paragraph,
    add_right_paragraph,
    add_paragraph_indent,
    add_wrapped_paragraph,
    add_page_break,
)


def doc_chapter1(
    sec11_p1, sec11_p2, sec11_p3,
    purpose_count, purpose_1, purpose_2, purpose_3,
    hypo_paragraph, hypo_items, scope_data,
    para_premise_str, premise_data, def_items, benefit_items
):
    # 🔹 ใช้ฟังก์ชัน doc_setup() เพื่อสร้างเอกสารพร้อมตั้งค่าหน้ากระดาษและฟอนต์
    doc = doc_setup()

    # ---------------------------------------------------------------------
    # 1.0 หัวข้อหลักของบท
    # ---------------------------------------------------------------------
    add_center_paragraph(doc, "บทที่ 1", bold=True, font_size=20)
    add_center_paragraph(doc, "บทนำ\n", bold=True, font_size=20)

    # ---------------------------------------------------------------------
    # 1.1 ความเป็นมาและความสำคัญของปัญหา
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.1  ความเป็นมาและความสำคัญของปัญหา", bold=True)
    add_wrapped_paragraph(doc, sec11_p1, n=85, disth=True, custom_tap=0.8)
    add_wrapped_paragraph(doc, sec11_p2, n=85, disth=True, custom_tap=0.8)
    add_wrapped_paragraph(doc, sec11_p3, n=85, disth=True, custom_tap=0.8)
    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 1.2 วัตถุประสงค์
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.2  วัตถุประสงค์", bold=True)
    add_paragraph_indent(doc, f"1.2.1  {purpose_1}", custom_tap=0.8)
    add_paragraph_indent(doc, f"1.2.2  {purpose_2}", custom_tap=0.8)
    add_paragraph_indent(doc, f"1.2.3  {purpose_3}", custom_tap=0.8)
    add_page_break(doc, 1.5)

    # ---------------------------------------------------------------------
    # 1.3 สมมติฐาน
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.3  สมมติฐาน", bold=True)
    add_wrapped_paragraph(doc, hypo_paragraph, n=85, disth=True, custom_tap=0.8)
    for i, item in enumerate(hypo_items, start=1):
        if isinstance(item, dict):
            txt = (item.get("main") or "").strip()
        else:
            txt = (item or "").strip()
        if txt:
            add_paragraph_indent(doc, f"1.3.{i}  {txt}", custom_tap=0.8)

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 1.4 ขอบเขตของโครงงาน
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.4  ขอบเขตการทำโครงงาน", bold=True)
    for i, item in enumerate(scope_data, start=1):
        main = item.get('main', '').strip()
        if main:
            add_paragraph_indent(doc, f"1.4.{i}  {main}")
        for j, sub in enumerate(item.get('subs', []), start=1):
            sub = sub.strip()
            if sub:
                add_paragraph_indent(doc, f"\t1.4.{i}.{j}  {sub}")

    add_page_break(doc, 1.5)

    # ---------------------------------------------------------------------
    # 1.5 ข้อตกลงเบื้องต้น
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.5  ข้อตกลงเบื้องต้น", bold=True)
    add_wrapped_paragraph(doc, para_premise_str, n=85, disth=True, custom_tap=0.8)
    for i, item in enumerate(premise_data, start=1):
        main = item.get('main', '').strip()
        if main:
            add_paragraph_indent(doc, f"1.5.{i}  {main}")
        for j, sub in enumerate(item.get('subs', []), start=1):
            sub = sub.strip()
            if sub:
                add_paragraph_indent(doc, f"\t1.5.{i}.{j}  {sub}")
    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 1.6 นิยามศัพท์เฉพาะ
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.6  นิยามศัพท์เฉพาะ", bold=True)
    for i, item in enumerate(def_items, start=1):
        if isinstance(item, dict):
            txt = (item.get("main") or "").strip()
        else:
            txt = (item or "").strip()
        if txt:
            add_paragraph_indent(doc, f"1.6.{i}  {txt}", custom_tap=0.8)

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 1.7 ประโยชน์ที่คาดว่าจะได้รับ
    # ---------------------------------------------------------------------
    add_left_paragraph(doc, "1.7  ประโยชน์ที่คาดว่าจะได้รับ", bold=True)
    for i, item in enumerate(benefit_items, start=1):
        if isinstance(item, dict):
            txt = (item.get("main") or "").strip()
        else:
            txt = (item or "").strip()
        if txt:
            add_paragraph_indent(doc, f"1.7.{i}  {txt}",custom_tap=0.8)
            
            
            
            
            
            
    return doc


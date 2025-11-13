from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from django.http import HttpResponse

def format_authors(authors_list, lang='th'):
    """Helper function to format a list of authors into a string."""
    if not authors_list:
        return 'N.p.' if lang == 'en' else 'ม.ป.ผู้.' # No author
    
    if len(authors_list) == 1:
        return authors_list[0]
    
    if lang == 'en':
        return ', '.join(authors_list[:-1]) + f', {authors_list[-1]}'
    else: # lang == 'th'
        return ', '.join(authors_list[:-1]) + f', {authors_list[-1]}'

def doc_refer(references):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    section = doc.sections[0]
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    head_1 = doc.add_paragraph("บรรณานุกรม")
    head_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head_1.runs[0].bold = True

    for ref in references:
        ref_type = ref.get('ref_type')
        ref_count = ref.get('ref_count', '')
        lang = ref.get('language', 'th')
        
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)

        text = ""
        
        if ref_type == '1':
            authors_str = format_authors(ref.get('authors', []), lang)
            title = ref.get('title', 'N.t.' if lang == 'en' else 'ม.ป.ร.')
            url = ref.get('url', '')
            access_date = ref.get('access_date', '')
            if lang == 'en':
                text = f"{authors_str}. {title}. Available at: {url}. Accessed {access_date}."
            else:
                text = f"{authors_str}. [อินเทอร์เน็ต]. {title}. จาก: {url}. [สืบค้นเมื่อวันที่ {access_date}]."
        
        elif ref_type == '2':
            authors_str = format_authors(ref.get('authors', []), lang)
            title = ref.get('title', '')
            print_count = ref.get('print_count', '')
            city_print = ref.get('city_print', '[place unknown] or [ม.ป.ท.]')
            publisher = ref.get('publisher', '[publisher unknown] or [ม.ป.พ.]')
            y_print = ref.get('y_print', '[date unknown] or [ม.ป.ป.]')
            edition_text = f'{print_count}st/nd/rd/th ed. ' if print_count else ''
            if lang == 'en':
                 text = f'{authors_str}. {title}. {edition_text}{city_print}: {publisher}; {y_print}.'
            else:
                 text = f'{authors_str}. {title}. พิมพ์ครั้งที่ {print_count}. {city_print}: {publisher}; {y_print}.'

        elif ref_type == '3':
            article_author = ref.get('article_author', '')
            article_title = ref.get('article_title', '')
            editor = ref.get('editor', '')
            book_title = ref.get('book_title', '')
            city_print = ref.get('city_print', '')
            publisher = ref.get('publisher', '')
            y_print = ref.get('y_print', '')
            pages = ref.get('pages', '')
            if lang == 'en':
                text = f'{article_author}. {article_title}. In: {editor}, editor. {book_title}. {city_print}: {publisher}; {y_print}. p. {pages}.'
            else:
                text = f'{article_author}. {article_title}. ใน: {editor}, บรรณาธิการ. {book_title}. {city_print}: {publisher}; {y_print}. หน้า {pages}.'
        
        elif ref_type == '4':
            author = ref.get('author', '')
            title = ref.get('title', '')
            format_type = ref.get('format', '')
            city_prod = ref.get('city_prod', '')
            publisher = ref.get('publisher', '')
            y_prod = ref.get('y_prod', '')
            if lang == 'en':
                text = f'{author}. {title} [{format_type}]. {city_prod}: {publisher}; {y_prod}.'
            else:
                text = f'{author}. {title} [{format_type}]. {city_prod}: {publisher}; {y_prod}.'
        
        elif ref_type == '5':  
            author = ref.get('author', '')
            article_title = ref.get('article_title', '')
            journal_name = ref.get('journal_name', '')
            pub_date = ref.get('pub_date', '')
            vol_issue = ref.get('volume_issue', '')
            pages = ref.get('pages', '')
            if lang == 'en':
                text = f'{author}. {article_title}. {journal_name} {pub_date};{vol_issue}:{pages}.'
            else:
                text = f'{author}. {article_title}. {journal_name} {pub_date};{vol_issue}:{pages}.'
        
        elif ref_type == '6':
            author = ref.get('author', '')
            article_title = ref.get('article_title', '')
            journal_name = ref.get('journal_name', '')
            resource_type = ref.get('resource_type', '[serial online]')
            db_update_date = ref.get('db_update_date', '')
            access_date = ref.get('access_date', '')
            url = ref.get('url', '')
            if lang == 'en':
                text = f'{author}. {article_title}. {journal_name} [{resource_type}] [updated {db_update_date}; cited {access_date}]. Available from: {url}'
            else:
                text = f'{author}. {article_title}. {journal_name} [{resource_type}] [ปรับปรุงเมื่อ {db_update_date}; อ้างเมื่อ {access_date}]. เข้าถึงได้จาก: {url}'
        
        elif ref_type == '7':
            editor = ref.get('editor', '')
            title = ref.get('title', '')
            conf_name = ref.get('conference_name', '')
            conf_date = ref.get('conference_date', '')
            conf_loc = ref.get('conference_location', '')
            city_print = ref.get('city_print', '')
            publisher = ref.get('publisher', '')
            y_print = ref.get('y_print', '')
            if lang == 'en':
                text = f'{editor}, editor. {title}. Proceedings of the {conf_name}; {conf_date}; {conf_loc}. {city_print}: {publisher}; {y_print}.'
            else:
                text = f'{editor}, บรรณาธิการ. {title}. การประชุม {conf_name}; {conf_date}; {conf_loc}. {city_print}: {publisher}; {y_print}.'
        
        elif ref_type == '8':
            presenter = ref.get('presenter', '')
            pres_title = ref.get('presentation_title', '')
            editor = ref.get('editor', '')
            conf_name = ref.get('conference_name', '')
            conf_date = ref.get('conference_date', '')
            conf_loc = ref.get('conference_location', '')
            city_print = ref.get('city_print', '')
            publisher = ref.get('publisher', '')
            y_print = ref.get('y_print', '')
            page = ref.get('page', '')
            if lang == 'en':
                text = f'{presenter}. {pres_title}. In: {editor}, editor. {conf_name}; {conf_date}; {conf_loc}. {city_print}: {publisher}; {y_print}. p. {page}.'
            else:
                text = f'{presenter}. {pres_title}. ใน: {editor}, บรรณาธิการ. {conf_name}; {conf_date}; {conf_loc}. {city_print}: {publisher}; {y_print}. หน้า {page}.'

        elif ref_type == '9':
            author = ref.get('author', '')
            article_title = ref.get('article_title', '')
            newspaper_name = ref.get('newspaper_name', '')
            pub_date = ref.get('pub_date', '')
            section = ref.get('section', '')
            page = ref.get('page', '')
            if lang == 'en':
                text = f'{author}. {article_title}. {newspaper_name} {pub_date};{section}:{page}.'
            else:
                text = f'{author}. {article_title}. {newspaper_name} {pub_date};{section}:{page}.'

        if text:
            p.add_run(f'[{ref_count}] ').bold = False
            p.add_run(text)
    
    return doc
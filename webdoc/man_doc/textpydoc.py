from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION

# doc = Document()

# doc.add_paragraph("บทที่ 1: บทนำ")

# # ขึ้นหน้าแบบเร็ว
# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# doc.add_paragraph("บทที่ 2: วัตถุประสงค์")

# # แยก section พร้อมขึ้นหน้าใหม่
# doc.add_section(WD_SECTION.NEW_PAGE)
# doc.add_paragraph("บทที่ 3: ขอบเขต")

# doc.save("report.docx")


from docx import Document

# รับข้อความจากคีย์บอร์ด
text = input("กรุณาพิมพ์ข้อความ: ")

# นับจำนวนตัวอักษร (ไม่รวมช่องว่าง)
num_chars = len(text.replace(" ", ""))
print(f"จำนวนตัวอักษร (ไม่นับช่องว่าง): {num_chars}")

# สร้างไฟล์ docx และเพิ่มข้อความลงไป



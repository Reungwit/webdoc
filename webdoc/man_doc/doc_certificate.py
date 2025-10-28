# man_doc/doc_certificate.py
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# Doc_cert เอา _ ออก
# -------------------- Page Border --------------------
def add_page_border(doc, line_size_halfpt=12, space_pt=8, color="000000"):
    """
    ใส่กรอบรอบหน้ากระดาษ (Page Border) แบบเส้นตรง
    line_size_halfpt: ความหนา (หน่วย half-point) เช่น 12 = 6pt
    space_pt: ระยะห่างจาก margin (เพิ่ม 8pt กันชนโลโก้/ลายเซ็น)
    """
    section = doc.sections[0]
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')

    def _edge(tag):
        e = OxmlElement(tag)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(line_size_halfpt))
        e.set(qn('w:space'), str(space_pt))
        e.set(qn('w:color'), color)
        return e

    pgBorders.append(_edge('w:top'))
    pgBorders.append(_edge('w:bottom'))
    pgBorders.append(_edge('w:left'))
    pgBorders.append(_edge('w:right'))
    section._sectPr.append(pgBorders)

# -------------------- helpers --------------------
def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "TH SarabunPSK"
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    st.element.rPr.rFonts.set(qn("w:ascii"), "TH SarabunPSK")
    st.element.rPr.rFonts.set(qn("w:hAnsi"), "TH SarabunPSK")
    st.font.size = Pt(16)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = 1.0

def center(doc, txt, bold=False, spacing_before_pt=0, spacing_after_pt=0):
    p = doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(spacing_before_pt)
    p.paragraph_format.space_after = Pt(spacing_after_pt)
    if p.runs:
        p.runs[0].bold = bold
    return p

def left(doc, txt, spacing_before_pt=0, spacing_after_pt=0):
    p = doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(spacing_before_pt)
    p.paragraph_format.space_after = Pt(spacing_after_pt)
    return p

def right(doc, txt, spacing_before_pt=0, spacing_after_pt=0, indent_spaces=2):
    p = doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(spacing_before_pt)
    p.paragraph_format.space_after = Pt(spacing_after_pt)
    # ขยับเข้ามาจากขอบขวาประมาณ 2 เคาะ
    p.paragraph_format.right_indent = Inches(0.22) if indent_spaces else Inches(0)
    return p

def hanging_label_paragraph(doc, label, text, label_width_cm=3.2, spacing_after_pt=0):
    """
    ย่อหน้าที่มี label (เช่น 'เรื่อง', 'โดย') แบบ hanging indent
    หลัง label เว้น 2 เคาะ แล้วข้อความทั้งหมดชิดคอลัมน์เดียวกัน
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(label_width_cm)
    pf.first_line_indent = -Cm(label_width_cm)
    pf.space_after = Pt(spacing_after_pt)
    run = p.add_run(f"{label}  ")  # ✅ เว้น 2 เคาะ
    run.bold = True
    p.add_run(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def hanging_continuation(doc, text, label_width_cm=3.2, spacing_after_pt=0):
    """
    ย่อหน้าต่อเนื่อง (ไม่มี label) ใช้ left_indent เดียวกับเนื้อหาหลัง label
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(label_width_cm)
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(spacing_after_pt)
    p.add_run(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def underline_line(length=40):
    return "_" * length

# ----------------------------------------------------
#                  MAIN GENERATOR
# ----------------------------------------------------

def doc_certificate(topic, author1, author2, comm_dean, prathan, comm_first, comm_sec):
    doc = Document()
    set_base_style(doc)

    # margins: บน/ซ้าย 1.5", ขวา/ล่าง 1"
    sec = doc.sections[0]
    sec.top_margin = Inches(1.5)
    sec.left_margin = Inches(1.5)
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    # page border
    add_page_border(doc)

    # ---------- LOGO ----------
    logo_path = os.path.join('static', 'img', 'kmutnb_logo_cover.png')
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.35))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    center(doc, "ใบรับรองปริญญานิพนธ์", bold=True, spacing_after_pt=2)
    center(doc, "คณะเทคโนโลยีและการจัดการอุตสาหกรรม", bold=True, spacing_after_pt=2)
    center(doc, "มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ", bold=True, spacing_after_pt=2)

    # ---------- เรื่อง / โดย (Hanging Indent) ----------
    LABEL_W = 0.9
    hanging_label_paragraph(doc, "เรื่อง", topic or "", label_width_cm=LABEL_W)
    hanging_label_paragraph(doc, "โดย", author1 or "", label_width_cm=LABEL_W)
    if (author2 or "").strip():
        hanging_continuation(doc, author2, label_width_cm=LABEL_W)
    doc.add_paragraph()
    # ---------- ข้อความหลักใต้เรื่อง ----------
    hanging_continuation(doc, "ได้รับอนุมัติให้นับเป็นส่วนหนึ่งของการศึกษาตาม", label_width_cm=LABEL_W)
    hanging_continuation(doc, "หลักสูตรอุตสาหกรรมศาสตรบัณฑิต สาขาวิชาเทคโนโลยีสารสนเทศ",
                          label_width_cm=LABEL_W, spacing_after_pt=12)
    doc.add_paragraph()
    # ---------- คณบดี (ชิดขวา + เว้นด้วย \t จากเส้น) ----------
    right(doc, f"{underline_line(40)}คณบดี", spacing_after_pt=2, indent_spaces=2)
    left(doc, f"\t\t\t\t\t({comm_dean.strip()})", spacing_after_pt=8)

    # ---------- คณะกรรมการ ----------
    left(doc, "คณะกรรมการสอบปริญญานิพนธ์", spacing_after_pt=6)

    # ประธาน
    left(doc, f"{underline_line(40)}ประธานกรรมการ", spacing_after_pt=1)
    left(doc, f"\t({prathan.strip()})" if (prathan or "").strip() else "", spacing_after_pt=8)

    # กรรมการ 1
    left(doc, f"{underline_line(40)}กรรมการ", spacing_after_pt=1)
    left(doc, f"\t({comm_first.strip()})" if (comm_first or "").strip() else "", spacing_after_pt=8)

    # กรรมการ 2 (ถ้ามี)
    if (comm_sec or "").strip():
        left(doc, f"{underline_line(40)}กรรมการ", spacing_after_pt=1)
        left(doc, f"\t({comm_sec.strip()})", spacing_after_pt=8)
    return doc

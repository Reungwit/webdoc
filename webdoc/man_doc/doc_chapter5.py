# doc_chapter5.py

from typing import List, Dict, Any, Optional
from docx import Document

# ✅ ดึง “ฟังก์ชันกลาง”
from man_doc.doc_function import *

DEFAULT_TITLES = [
    "บทนำ",                    # (ไม่ใส่เลขหัวข้อ)
    "สรุปผลการดำเนินงาน",      # 5.1
    "อภิปรายผล",                # 5.2
    "ข้อเสนอแนะ",               # 5.3
]

def doc_chapter5(
    sections_any: Any,                 # โครงจาก DB/UI (ผสมได้)
    *,
    media_root: str = "",
) -> Document:
    """
    เรนเดอร์บทที่ 5:
      - บรรทัดหัวบทกึ่งกลาง
      - 'บทนำ' ไม่มีเลขหัวข้อ (เพียงย่อหน้าเนื้อหา)
      - 5.1 เป็นต้นไป = หัวข้อหลัก (สรุปผลฯ, อภิปรายผล, ข้อเสนอแนะ)
      - หัวข้อย่อยลึกลงไปใช้เลข 5.x.y, 5.x.y.z อัตโนมัติ
      - แคปชันรูป = 'ภาพที่ 5-n ...'
    """

    # ---------- เตรียมเอกสาร + normalize โครงสร้าง ----------
    doc = doc_setup()

    # หัวกระดาษบท
    add_center_paragraph(doc, "บทที่ 5", bold=True, font_size=18)
    add_center_paragraph(doc, "สรุปผลการวิจัย อภิปรายผล และข้อเสนอแนะ", bold=True, font_size=18)

    # ปรับให้โครงสร้างปลอดภัยสำหรับเรนเดอร์
    secs = sections_doc_safe(
        sections_any,
        default_titles=DEFAULT_TITLES,
        first_section_mode="paragraphs",   # บทนำเก็บเป็น paragraphs
    )

    # ---------- 0) บทนำ (ไม่มีเลขหัวข้อ) ----------
    intro = secs[0]
    paras = intro.get("paragraphs") or []
    for s in paras:
        # ย่อหน้าแนวบรรยาย
        add_wrapped_paragraph(doc, str(s), n=120, disth=True, tap=True)

    # ---------- 1) 5.1–5.3 หัวข้อหลัก + รายการย่อย ----------
    chapter_no = 5
    pic_counter = [0]
    seen_pics = set()

    # index: 1..3  → หมายเลขหัวข้อ: 5.1, 5.2, 5.3
    for idx in range(1, min(4, len(secs))):
        sec = secs[idx]
        title = (sec.get("title") or "").strip()
        body  = (sec.get("body") or "").strip()

        # หัวข้อระดับ 1: 5.{idx}
        title_no = f"{chapter_no}.{idx}"
        add_section_heading_level1_style_custom(doc, title_no, title)

        # ย่อหน้าเนื้อหาใต้หัวข้อ
        if body:
            add_body_paragraph_style_1(doc, body)

        # เดินรายการย่อย (หัวข้อ 2 ขึ้นไป → 5.{idx}.1, 5.{idx}.2, …)
        # heading_func = สไตล์หัวข้อระดับ 2+
        # body_func    = สไตล์ย่อหน้า
        walk_item_tree(
            doc,
            section_no=title_no,                 # เริ่มเลขจาก 5.{idx}
            nodes=sec.get("mains") or [],
            chapter_no=chapter_no,               # แคปชันรูปเป็น 5-n
            media_root=media_root,
            pic_counter=pic_counter,
            seen_pics=seen_pics,
            heading_func=add_section_heading_level2_plus_style_custom,
            body_func=add_body_paragraph_style_1,
        )

    return doc


def add_section_heading_level1_style_custom(doc: Document, title_no: str, title: str) -> None:
    """
    (ฟังก์ชันกลาง) สไตล์หัวข้อระดับ 1: หนา, ไม่ย่อหน้า, ห่าง 6pt
    """
    text = two_spaces_join(t(title_no), t(title)) # (ต้อง import two_spaces_join)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = True
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    
def add_section_heading_level2_plus_style_custom(doc: Document, title_no: str, title: str) -> None:
    """
    (ฟังก์ชันกลาง) สไตล์หัวข้อระดับ 2+: หนา, ย่อ 0.75, ห่าง 3pt
    """
    text = two_spaces_join(t(title_no), t(title)) # (ต้อง import two_spaces_join)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = False
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.75)
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
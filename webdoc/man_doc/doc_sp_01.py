from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK # เหมือนกด Ctrl + Enter ใน Word
from pythainlp.tokenize import word_tokenize #ใช้ตัดคำ
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document as DocxDocument
from man_doc.doc_function import *
def doc_sp_01(name_pro_th, name_pro_en, authors,case_stu,
              term,school_y,adviser,co_advisor,strategic,
              plan,key_result,bg_and_sig_para1,bg_and_sig_para2,bg_and_sig_para3,
              purpose_1,purpose_2,purpose_3,scope_data):
    

    doc = Document()
    
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    # style.font.bold = True
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    # ตั้งค่าขอบกระดาษ
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8858)
    section.left_margin = Inches(1.248)
    section.right_margin = Inches(0.748)
    add_page_number(section)

    
    
    head_1 = doc.add_paragraph("ทก.01")
    head_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head_1.runs[0].bold = True
    
    
    add_center_paragraph(doc, "แบบเสนอโครงงานพิเศษ (ปริญญานิพนธ์)", bold=True)
    add_center_paragraph(doc, "ภาควิชาเทคโนโลยีสารสนเทศ", bold=True)
    add_center_paragraph(doc, "คณะเทคโนโลยีและการจัดการอุตสาหกรรม" , bold=True)
    add_left_paragraph(doc , "1.ข้อมูลขั้นต้นของโครงงาน" , bold=True)
    p = doc.add_paragraph()
    p.add_run("\t1.1 ชื่อโครงงาน (ภาษาไทย)\t").bold = True
    p.add_run(name_pro_th).bold = False
    add_paragraph_indent(doc, f"\t\t(ภาษาอังกฤษ)\t{name_pro_en}")
    p = doc.add_paragraph()
    p.add_run(f"\t\tกรณีศึกษา").bold = True
    p.add_run(f" (ถ้ามี)\t{case_stu}\n").bold = False
    add_paragraph_indent(doc ,"1.2 ชื่อนักศึกษาผู้ทำโครงงาน",bold=True)
    
    # รายชื่อแต่ละคน
    for idx, name in enumerate(authors, start=1):
        add_paragraph_indent(doc,f"\t{idx}. {name}",bold=False)
    doc.add_paragraph("")
    add_paragraph_indent(doc ,"\tหลักสูตรอุตสาหกรรมศาสตรบัณฑิต   ภาควิชาเทคโนโลยีสารสนเทศ (ต่อเนื่อง)",bold=False)
    add_paragraph_indent(doc ,f"\tภาคเรียนที่\t{term}\tปีการศึกษา {school_y} ",bold=False)
    doc.add_paragraph("")
    add_paragraph_indent(doc , "1.3 ชื่ออาจารย์ที่ปรึกษา ", bold=True)
    add_paragraph_indent(doc , f"\t\t{adviser}", bold=False)
    add_paragraph_indent(doc , "     อาจารย์ที่ปรึกษาร่วม ", bold=True)
    add_paragraph_indent(doc , f"\t\t{co_advisor}\n", bold=False)
    add_paragraph_indent(doc , "1.4 ความสอดคล้องกับแผนงานด้านวิทยศาสตร์ วิจัยละนวัตกรรม", bold=True)
    
    
    p = doc.add_paragraph()
    p.add_run(f"\tยุทธศาสตร์ที่ ").bold = True
    add_wrapped_paragraph(p, text=f"{strategic}", n=85 )
    # p.add_run(strategic)
    # p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # set_thai_distributed(p)

    p = doc.add_paragraph()
    p.add_run(f"\tแผนงานที่ (P) ").bold = True
    p.add_run(plan)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # set_thai_distributed(p)

    p = doc.add_paragraph()
    p.add_run(f"\tผลลัพธ์ที่สำคัญ (key result) ").bold = True
    add_wrapped_paragraph(p, text=f"{key_result}", n=93 ,disth=True )
    
    
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    # doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_left_paragraph(doc , "2. รายละเอียดโครงงาน" , bold=True)
    add_paragraph_indent(doc, "2.1 ความเป็นมาและความสำคัญของปัญหา" , bold=True)
    add_wrapped_paragraph(doc, text=f"{bg_and_sig_para1}", n=93 ,disth=True ,extap=True)
    add_wrapped_paragraph(doc, text=f"{bg_and_sig_para2}", n=93 ,disth=True ,extap=True)
    add_wrapped_paragraph(doc, text=f"{bg_and_sig_para3}", n=93 ,disth=True ,extap=True)
    # p = doc.add_paragraph() #เรียกฟังก์ชันตัดคำใช้กับ addrun
    # add_wrapped_paragraph(p, text=f"{bg_and_sig_para1}", n=90)
    
    doc.add_paragraph("")
    add_left_paragraph(doc , "2.2 วัตถุประสงค์" , bold=True)
    add_paragraph_indent(doc, f"2.2.1 {purpose_1}")
    add_paragraph_indent(doc, f"2.2.2 {purpose_2}")
    add_paragraph_indent(doc, f"2.2.3 {purpose_3}")
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    p = doc.add_paragraph()
    p.add_run("2.3 ขอบเขตการทำโครงงาน ").bold = True
    p.add_run("(Scope of Special Project)")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # ========== ส่วนขอบเขต ==========
    # doc.add_heading('2.3 ขอบเขตการทำโครงงาน (Scope of Special Project)', level=2)
    for i, item in enumerate(scope_data, start=1):
        main = item.get('main', '').strip()
        if main:
            add_paragraph_indent(doc, f"2.3.{i} {main}")  # ✅ ไม่มีเลขอัตโนมัติ

        for j, sub in enumerate(item.get('subs', []), start=1):
            sub = sub.strip()
            if sub:
                add_paragraph_indent(doc, f"\t2.3.{i}.{j} {sub}")  # ✅ ไม่มี bullet

    
    
    
    
    
    return doc


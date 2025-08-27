# doc_intro.py
# --------------------------------------------------
# เอกสาร Word สำหรับ:
#   (1) บทคัดย่อภาษาไทย
#   (2) Abstract ภาษาอังกฤษ
#   (3) กิตติกรรมประกาศ (ไม่มีลายเซ็น)
#
# ไฮไลต์:
#   - ใช้ Tab Stops ให้ ":" ตรงคอลัมน์เดียวกันทุกหัวข้อ
#   - บรรทัดที่ 2 ของหัวข้อ (ผู้จัดทำ/ชื่อเรื่อง/สาขา/ที่ปรึกษา) ไม่ใส่ ":" และเยื้องให้ตรงคอลัมน์ค่า
#   - สาขาวิชาเติมบรรทัดมหาวิทยาลัยให้อัตโนมัติ (TH/EN)
#   - แสดงจำนวนหน้า (ไทย/อังกฤษ) ชิดขวาหลังเนื้อหา
#   - คำสำคัญ/Keywords เว้น 1 บรรทัดจากเนื้อหา
#   - บรรทัดลายเซ็นวางที่ "ก้นหน้ากระดาษในส่วนเนื้อหา" (ไม่ใช่ footer) โดยใช้ framePr ยึดตำแหน่งกับหน้า
# --------------------------------------------------

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pythainlp.tokenize import word_tokenize

# ======= ค่าคงที่ตำแหน่ง Tab Stops (เซนติเมตร) =======
LABEL_TAB_CM = 3.5   # ตำแหน่งสิ้นสุดคอลัมน์ Label
COLON_TAB_CM = 4.0   # ตำแหน่งเครื่องหมาย ":"


# ==================================================
# 🔹 Helper สำหรับการจัดย่อหน้า/ตัวอักษร
# ==================================================

def set_thai_distributed(paragraph):
    """จัดบรรทัดภาษาไทยให้เต็มบรรทัด (thaiDistribute) เพื่อการกระจายคำที่สวยกว่า justify ธรรมดา"""
    p_pr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "thaiDistribute")
    p_pr.append(jc)

def add_center_paragraph(doc, text, bold=True):
    """เพิ่มพารากราฟกึ่งกลาง (เช่น หัวเรื่อง 'บทคัดย่อ', 'Abstract')"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = bold
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(16)
    return p

def add_wrapped_paragraph(doc, text, n=85, indent=True):
    """
    เพิ่มย่อหน้าเนื้อหาแบบตัดคำอัตโนมัติ (ไทย)
    - n: จำนวนอักขระประมาณต่อบรรทัด เพื่อช่วยตัดบรรทัด
    - indent: ย่อบรรทัดแรก 1.27 ซม.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)

    words = word_tokenize(text or "", engine="newmm")
    buf = ""
    for w in words:
        if len(buf + w) <= n:
            buf += w
        else:
            p.add_run(buf.strip())
            p.add_run().add_break()
            buf = w
    if buf:
        p.add_run(buf.strip())

    for r in p.runs:
        r.font.name = "TH SarabunPSK"
        r.font.size = Pt(16)

    set_thai_distributed(p)
    return p

def add_left_paragraph(doc, text, space_before_pt=None):
    """เพิ่มบรรทัดชิดซ้ายทั่วไป (กำหนดเว้นระยะก่อนบรรทัดได้)"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    r = p.runs[0]
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(16)
    return p

def add_right_paragraph(doc, text, space_before=12):
    """เพิ่มบรรทัดชิดขวา (ใช้แสดงจำนวนหน้า)"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(space_before)
    r = p.runs[0]
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(16)
    return p


# ==================================================
# 🔹 Helper จัด Tab Stops ให้ ":" ตรงคอลัมน์
# ==================================================

def _set_tabstops(p):
    """ตั้งค่า Tab Stops ให้ทุกบรรทัดของหัวข้อวาง Label, ':' และค่า ได้ตรงตำแหน่งเดียวกัน"""
    ts = p.paragraph_format.tab_stops
    for t in list(ts):
        ts.remove(t)
    ts.add_tab_stop(Cm(LABEL_TAB_CM))
    ts.add_tab_stop(Cm(COLON_TAB_CM))

def _add_info_line(doc, label, first_value_line):
    """บรรทัดแรกของหัวข้อ: 'Label<TAB>:<TAB>Value'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_tabstops(p)
    r = p.add_run(f"{label}\t:\t{first_value_line}")
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(16)

def _add_cont_line(doc, value_line):
    """บรรทัดถัดไปของหัวข้อ: ไม่มี ':' และเยื้องไปคอลัมน์ค่าโดยกดแท็บ 2 ครั้ง"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_tabstops(p)
    r = p.add_run(f"\t\t{value_line}")
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(16)

def add_info_block_with_tabs(doc, pairs):
    """
    วาดบล็อคหัวข้อมูลด้วย Tab Stops
    pairs: [(label: str, value_lines: list[str]), ...]
      - value_lines[0] = บรรทัดแรก (แสดง ':')
      - value_lines[1..] = บรรทัดต่อ ๆ ไป (ไม่แสดง ':')
    """
    for label, lines in pairs:
        lines = lines or [""]
        _add_info_line(doc, label, lines[0].strip())
        for cont in lines[1:]:
            _add_cont_line(doc, cont.strip())


# ==================================================
# 🔹 วาง "ลายเซ็น" ที่ก้นหน้ากระดาษ (ส่วนเนื้อหา) ด้วย framePr
# ==================================================

def _emu_from_pt(pt_val: float) -> int:
    """แปลงจุด (pt) เป็น EMU (1 pt = 12700 EMU)"""
    return int(round(pt_val * 12700))

def _twips_from_pt(pt_val: float) -> int:
    """แปลงจุด (pt) เป็น twips (1 pt = 20 twips)"""
    return int(round(pt_val * 20))

def add_signature_at_page_bottom(doc, section, caption: str, underline_len: int = 60, bottom_gap_pt: int = 6):
    """
    เพิ่มย่อหน้า "เส้นสำหรับเซ็น + ข้อความ" ไว้ก้นหน้ากระดาษ (ภายใน body ไม่ใช่ footer)
    โดยกำหนดตำแหน่งย่อหน้าด้วย w:framePr ยึดกับหน้า/ขอบกระดาษ

    พารามิเตอร์:
      - caption         : ข้อความต่อท้ายเส้น (เช่น 'อาจารย์ที่ปรึกษาปริญญานิพนธ์' / 'Project Advisor')
      - underline_len   : จำนวน "_" เพื่อความยาวเส้น
      - bottom_gap_pt   : ระยะเผื่อจากขอบล่าง (pt) เพื่อไม่ให้ชนขอบพิมพ์ (เช่น 6–12pt)

    หลักการ:
      - วาง vAnchor='page'  และ hAnchor='margin'
      - คำนวณ y (ตำแหน่งแนวตั้ง) เป็น: page_height - bottom_margin - bottom_gap - สูงโดยประมาณของบรรทัด
      - ใส่ลงในพารากราฟสุดท้ายของ section (หลังเนื้อหาส่วนไทย/อังกฤษ)
    """
    # --- คำนวณพิกัด y เป็น twips ---
    page_height_pt = section.page_height.pt
    bottom_margin_pt = section.bottom_margin.pt

    # สูงโดยประมาณของบรรทัดลายเซ็น (ฟอนต์ 16pt line spacing 1.0)
    line_pt = 16.0

    # y จากขอบบนหน้า → ไปยังตำแหน่งใกล้ bottom margin
    y_pt = page_height_pt - bottom_margin_pt - bottom_gap_pt - line_pt
    y_tw = _twips_from_pt(y_pt)

    # --- สร้างย่อหน้าและปรับ framePr ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # เขียนเส้น + เว้นช่องว่างไม่ตัดบรรทัด + คำอธิบาย
    run1 = p.add_run("_" * underline_len + "\u00A0\u00A0")
    run1.font.name = "TH SarabunPSK"
    run1.font.size = Pt(16)

    run2 = p.add_run(caption)
    run2.font.name = "TH SarabunPSK"
    run2.font.size = Pt(16)

    # สร้าง w:framePr เพื่อยึดตำแหน่งย่อหน้านี้กับหน้า
    pPr = p._p.get_or_add_pPr()
    framePr = OxmlElement('w:framePr')

    # ยึดแนวตั้งกับ "page", แนวนอนกับ "margin"
    framePr.set(qn('w:vAnchor'), 'page')
    framePr.set(qn('w:hAnchor'), 'margin')

    # ตำแหน่ง y (twips) จากขอบบนของหน้า
    framePr.set(qn('w:y'), str(y_tw))

    # ล็อคตำแหน่ง (กัน Word ปรับเลื่อน)
    framePr.set(qn('w:anchorLock'), '1')

    # ไม่ให้ wrap ข้อความรอบ ๆ
    framePr.set(qn('w:wrap'), 'none')

    # แนวนอนให้กึ่งกลาง margin (x ไม่ระบุ ปล่อยให้จัดกลางด้วย alignment)
    # ถ้าอยาก fix x เพิ่มเติม (เช่นชิดขวา) สามารถ set w:x ได้เช่นกัน

    pPr.append(framePr)

    return p


# ==================================================
# 🔹 ส่วนหัวข้อมูล (ไทย/อังกฤษ)
# ==================================================

def add_header_info_th(document, data):
    """ประกอบหัวข้อมูลภาษาไทย พร้อมบรรทัด 2 อัตโนมัติของมหาวิทยาลัย"""
    author_lines = [data.get("author1_th", "") or ""]
    if data.get("author2_th"):
        author_lines.append(data.get("author2_th", ""))

    proj_lines = [data.get("project_name_th", "") or ""]
    if data.get("project_name_th_line2"):
        proj_lines.append(data.get("project_name_th_line2", ""))

    major_lines = [data.get("major_th", "") or "",
                   "มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ"]

    advisor_lines = [data.get("advisor_th", "") or ""]
    if data.get("coadvisor_th"):
        advisor_lines.append(data.get("coadvisor_th", ""))

    pairs = [
        ("ชื่อ", author_lines),
        ("ชื่อปริญญานิพนธ์", proj_lines),
        ("สาขาวิชา", major_lines),
        ("ที่ปรึกษาปริญญานิพนธ์", advisor_lines),
        ("ปีการศึกษา", [str(data.get("academic_year_th", "") or "")]),
    ]
    add_info_block_with_tabs(document, pairs)

def add_header_info_en(document, data):
    """ประกอบหัวข้อมูลภาษาอังกฤษ พร้อมบรรทัด 2 อัตโนมัติของมหาวิทยาลัย"""
    author_lines = [data.get("author1_en", "") or ""]
    if data.get("author2_en"):
        author_lines.append(data.get("author2_en", ""))

    proj_lines = [data.get("project_name_en", "") or ""]
    if data.get("project_name_en_line2"):
        proj_lines.append(data.get("project_name_en_line2", ""))

    major_lines = [data.get("major_en", "") or "",
                   "King Mongkut’s University of Technology North Bangkok"]

    advisor_lines = [data.get("advisor_en", "") or ""]
    if data.get("coadvisor_en"):
        advisor_lines.append(data.get("coadvisor_en", ""))

    pairs = [
        ("Name", author_lines),
        ("Project Title", proj_lines),
        ("Major Field", major_lines),
        ("Project Advisor", advisor_lines),
        ("Academic Year", [str(data.get("academic_year_en", "") or "")]),
    ]
    add_info_block_with_tabs(document, pairs)


# ==================================================
# 🔹 Main: ประกอบเอกสารทั้ง 3 ส่วน
# ==================================================

def doc_intro(data: dict):
    document = Document()

    # ----- ตั้งค่าสไตล์พื้นฐาน -----
    style = document.styles['Normal']
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ===== (1) หน้าบทคัดย่อภาษาไทย =====
    section_th = document.sections[0]
    section_th.top_margin = Cm(3.81)     # 1.5"
    section_th.bottom_margin = Cm(2.54)  # 1.0"
    section_th.left_margin = Cm(3.81)    # 1.5"
    section_th.right_margin = Cm(2.54)   # 1.0"

    add_header_info_th(document, data)
    add_center_paragraph(document, "บทคัดย่อ", True)
    add_wrapped_paragraph(document, data.get("abstract_th_para1", ""), n=87, indent=True)
    add_wrapped_paragraph(document, data.get("abstract_th_para2", ""), n=87, indent=True)

    # จำนวนหน้า (ไทย) ชิดขวา
    pages_th = str(data.get("total_pages", "") or "")
    add_right_paragraph(document, f"(ปริญญานิพนธ์มีจำนวนทั้งสิ้น {pages_th} หน้า)", space_before=6)

    # เว้น 1 บรรทัดแล้วคำสำคัญ
    add_left_paragraph(document, f"คำสำคัญ: {data.get('keyword_th', '')}", space_before_pt=16)

    # 🔻 ลายเซ็น: วาง "ในก้นหน้ากระดาษ" ของหน้าไทย (ไม่ใช่ footer)
    add_signature_at_page_bottom(document, section_th, "อาจารย์ที่ปรึกษาปริญญานิพนธ์",
                                 underline_len=51, bottom_gap_pt=8)

    # ===== (2) หน้า Abstract ภาษาอังกฤษ =====
    section_en = document.add_section(WD_SECTION.NEW_PAGE)
    section_en.top_margin = section_th.top_margin
    section_en.bottom_margin = section_th.bottom_margin
    section_en.left_margin = section_th.left_margin
    section_en.right_margin = section_th.right_margin

    add_header_info_en(document, data)
    add_center_paragraph(document, "Abstract", True)
    add_wrapped_paragraph(document, data.get("abstract_en_para1", ""), n=85, indent=True)
    add_wrapped_paragraph(document, data.get("abstract_en_para2", ""), n=85, indent=True)

    # จำนวนหน้า (อังกฤษ) ชิดขวา
    pages_en = str(data.get("total_pages", "") or "")
    add_right_paragraph(document, f"(Total {pages_en} Page)", space_before=6)

    # เว้น 1 บรรทัดแล้ว Keywords
    add_left_paragraph(document, f"Keywords: {data.get('keyword_en', '')}", space_before_pt=16)

    # 🔻 ลายเซ็น: วาง "ในก้นหน้ากระดาษ" ของหน้าอังกฤษ (ไม่ใช่ footer)
    add_signature_at_page_bottom(document, section_en, "Project Advisor",
                                 underline_len=62, bottom_gap_pt=8)

    # ===== (3) หน้า กิตติกรรมประกาศ (ไม่มีลายเซ็น) =====
    section_ack = document.add_section(WD_SECTION.NEW_PAGE)
    section_ack.top_margin = section_th.top_margin
    section_ack.bottom_margin = section_th.bottom_margin
    section_ack.left_margin = section_th.left_margin
    section_ack.right_margin = section_th.right_margin

    add_center_paragraph(document, "กิตติกรรมประกาศ", True)
    add_wrapped_paragraph(document, data.get("acknow_para1", ""), n=87, indent=True)
    add_wrapped_paragraph(document, data.get("acknow_para2", ""), n=87, indent=True)

    # ลงชื่อผู้จัดทำ (ชิดขวา) — หน้านี้ไม่มีลายเซ็นแบบขีดเส้น
    if data.get("acknow_name1"):
        add_right_paragraph(document, data.get("acknow_name1", ""), space_before=36)
    if data.get("acknow_name2"):
        add_right_paragraph(document, data.get("acknow_name2", ""), space_before=12)

    return document

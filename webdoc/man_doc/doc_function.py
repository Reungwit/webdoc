from __future__ import annotations
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK # ‡πÄ‡∏´‡∏°‡∏∑‡∏≠‡∏ô‡∏Å‡∏î Ctrl + Enter ‡πÉ‡∏ô Word
from docx.enum.section import WD_SECTION
from pythainlp.tokenize import word_tokenize #‡πÉ‡∏ä‡πâ‡∏ï‡∏±‡∏î‡∏Ñ‡∏≥
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document as DocxDocument
from typing import Any, Dict, Iterable, List, Tuple, Optional
import json
import os

def doc_setup():
    doc = Document()

    # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö‡∏ü‡∏≠‡∏ô‡∏ï‡πå‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
    style.font.size = Pt(16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ‚úÖ ‡πÉ‡∏ä‡πâ first section ‡πÅ‡∏ö‡∏ö‡∏õ‡∏•‡∏≠‡∏î‡∏†‡∏±‡∏¢ (‡∏Å‡∏±‡∏ô edge-case)
    section = doc.sections[0] if getattr(doc, "sections", None) and len(doc.sections) > 0 \
              else doc.add_section(WD_SECTION.NEW_PAGE)

    section.top_margin    = Inches(2.0)  # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î margin ‡∏´‡∏ô‡πâ‡∏≤‡πÅ‡∏£‡∏Å
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.5)
    section.right_margin  = Inches(1)

    return doc

def add_center_paragraph(doc, text, bold=False, font_size=16):
    p = doc.add_paragraph()
    r = p.add_run(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r.bold = bool(bold)
    r.font.size = Pt(font_size)
    return p

def add_left_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r.bold = bool(bold)
    return p

def add_right_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r.bold = bool(bold)
    return p


def add_paragraph_indent(doc, text, bold=False, custom_tap: float = 0.0):
    p = doc.add_paragraph()
    r = p.add_run(text or "")
    r.bold = bool(bold)
    p.paragraph_format.first_line_indent = Cm(custom_tap if custom_tap else 1.00)
    return p


def apply_rest_page_margin(doc: Document, top_inch: float = 1.5):
    """‡∏õ‡∏£‡∏±‡∏ö top margin ‡∏Ç‡∏≠‡∏á‡∏´‡∏ô‡πâ‡∏≤‡∏ñ‡∏±‡∏î‡πÑ‡∏õ (‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á ‡πÑ‡∏°‡πà‡∏Ç‡∏∂‡πâ‡∏ô‡∏´‡∏ô‡πâ‡∏≤‡πÉ‡∏´‡∏°‡πà)"""
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    base = doc.sections[0]
    sec.top_margin = Inches(top_inch)
    sec.bottom_margin = base.bottom_margin
    sec.left_margin = base.left_margin
    sec.right_margin = base.right_margin

def add_wrapped_paragraph(
    p_or_doc,
    text: str,
    n: int,
    disth: bool = False,
    extap: bool = False,
    tap: bool = False,
    custom_tap: float = 0.0,
    bold: bool = False,
):
    """
    ‡∏™‡∏£‡πâ‡∏≤‡∏á/‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ó‡∏µ‡πà‡∏ñ‡∏π‡∏Å‡∏ï‡∏±‡∏î‡∏Ñ‡∏≥‡∏•‡∏á‡πÉ‡∏ô paragraph ‡∏´‡∏£‡∏∑‡∏≠ document/cell
    - disth=True ‡πÉ‡∏ä‡πâ thaiDistribute (‡πÄ‡∏´‡∏°‡∏≤‡∏∞‡∏Å‡∏±‡∏ö‡∏†‡∏≤‡∏©‡∏≤‡πÑ‡∏ó‡∏¢)
    - bold=True ‡∏ó‡∏≥‡πÉ‡∏´‡πâ‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î‡πÉ‡∏ô‡∏¢‡πà‡∏≠‡∏´‡∏ô‡πâ‡∏≤‡∏ô‡∏µ‡πâ‡πÄ‡∏õ‡πá‡∏ô‡∏ï‡∏±‡∏ß‡∏´‡∏ô‡∏≤
    """

    def set_thai_distributed(paragraph):
        p_pr = paragraph._p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "thaiDistribute")
        p_pr.append(jc)

    # ‡πÄ‡∏ï‡∏£‡∏µ‡∏¢‡∏°‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏à‡∏≤‡∏Å‡∏Å‡∏≤‡∏£‡∏ï‡∏±‡∏î‡∏Ñ‡∏≥‡πÑ‡∏ó‡∏¢
    text = text or ""
    lines = []
    for paragraph in text.split("\n"):
        words = word_tokenize(paragraph.strip(), engine="newmm") #attacut / newmm
        line = ""
        for word in words:
            if len(line + word) <= n:
                line += word
            else:
                if line.strip():
                    lines.append(line.strip())
                line = word
        if line:
            lines.append(line.strip())

    # ‡∏ï‡∏£‡∏ß‡∏à‡∏ä‡∏ô‡∏¥‡∏î‡∏õ‡∏•‡∏≤‡∏¢‡∏ó‡∏≤‡∏á
    if isinstance(p_or_doc, (DocxDocument, _Cell)):
        p = p_or_doc.add_paragraph()
    elif isinstance(p_or_doc, Paragraph):
        p = p_or_doc
    else:
        raise TypeError("Argument must be Document, _Cell, or Paragraph")

    # ‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏•‡∏á paragraph
    for idx, l in enumerate(lines):
        # ‡∏ô‡∏±‡∏ö‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡πÅ‡∏ó‡πá‡∏ö‡∏ï‡πâ‡∏ô‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î ‡πÅ‡∏•‡πâ‡∏ß‡∏•‡∏ö‡∏≠‡∏≠‡∏Å‡∏Å‡πà‡∏≠‡∏ô‡∏û‡∏¥‡∏°‡∏û‡πå
        tab_count = len(l) - len(l.lstrip("\t"))
        l = l.lstrip("\t")

        # ‡πÉ‡∏™‡πà‡πÅ‡∏ó‡πá‡∏ö‡∏ï‡∏≤‡∏°‡∏à‡∏≥‡∏ô‡∏ß‡∏ô (‡πÑ‡∏°‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏´‡∏ô‡∏≤)
        for _ in range(tab_count):
            p.add_run().add_tab()

        # ‡πÉ‡∏™‡πà‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏à‡∏£‡∏¥‡∏á (‡∏´‡∏ô‡∏≤‡πÄ‡∏°‡∏∑‡πà‡∏≠ bold=True)
        if l.strip():
            r = p.add_run(l)
            if bold:
                r.bold = True

        # ‡∏Ç‡∏∂‡πâ‡∏ô‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡πÉ‡∏´‡∏°‡πà‡∏ñ‡πâ‡∏≤‡∏¢‡∏±‡∏á‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πà‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏™‡∏∏‡∏î‡∏ó‡πâ‡∏≤‡∏¢
        if idx < len(lines) - 1:
            p.add_run().add_break()

    # ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡∏Å‡∏≤‡∏£‡∏à‡∏±‡∏î‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö paragraph
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True

    if disth:
        set_thai_distributed(p)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ‡∏¢‡πà‡∏≠‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡πÅ‡∏£‡∏Å
    if extap:
        p.paragraph_format.first_line_indent = Cm(1.80)
    elif tap:
        p.paragraph_format.first_line_indent = Cm(1.00)
    elif custom_tap:
        p.paragraph_format.first_line_indent = Cm(custom_tap)

    return p

def add_page_break(doc, top_margin_inch=1.5):
    """
    ‡πÅ‡∏ó‡∏£‡∏Å Page Break ‡πÅ‡∏•‡πâ‡∏ß‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ç‡∏≠‡∏ö‡∏Å‡∏£‡∏∞‡∏î‡∏≤‡∏©‡∏î‡πâ‡∏≤‡∏ô‡∏ö‡∏ô‡∏Ç‡∏≠‡∏á‡∏´‡∏ô‡πâ‡∏≤‡∏ñ‡∏±‡∏î‡πÑ‡∏õ
    ‡πÄ‡∏£‡∏µ‡∏¢‡∏Å‡πÉ‡∏ä‡πâ add_page_break(doc)
        
    """
    # doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ‡πÄ‡∏û‡∏¥‡πà‡∏° section break ‡πÅ‡∏ö‡∏ö‡πÄ‡∏£‡∏¥‡πà‡∏°‡∏´‡∏ô‡πâ‡∏≤‡∏ñ‡∏±‡∏î‡πÑ‡∏õ
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    # ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤ margin ‡πÄ‡∏â‡∏û‡∏≤‡∏∞ section ‡∏ô‡∏µ‡πâ
    new_section.top_margin = Inches(top_margin_inch)
    
    


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1  # 0=left, 1=center, 2=right

    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

# -------------------------‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô ‡∏à‡∏±‡∏î‡∏Å‡∏≤‡∏£‡∏£‡∏π‡∏õ---------------------
# ===================== utils =====================

def t(v: Any) -> str:
    return (v or "").strip() if isinstance(v, str) else ""

def as_list(v: Any) -> list:
    return v if isinstance(v, list) else []

def resolve_image_path(p: Dict[str, Any], media_root: str) -> str | None:
    """‡∏Ñ‡∏∑‡∏ô absolute path ‡∏Ç‡∏≠‡∏á‡∏£‡∏π‡∏õ ‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö pic_path / pic_url (/media/...)"""
    rel = t(p.get("pic_path"))
    if rel:
        rel = rel.replace("\\", "/").lstrip("/")
        abs_path = os.path.join(media_root or "", rel)
        abs_path = abs_path.replace("\\", os.sep)
        if os.path.isfile(abs_path):
            return abs_path
    url = t(p.get("pic_url"))
    if url:
        url = url.replace("\\", "/").lstrip("/")
        if url.lower().startswith("media/"):
            url = url[6:]
        abs_path = os.path.join(media_root or "", url).replace("\\", os.sep)
        if os.path.isfile(abs_path):
            return abs_path
    return None

def two_spaces_join(a: str, b: str) -> str:
    """‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏°‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏î‡πâ‡∏ß‡∏¢ '  ' (2 ‡∏ä‡πà‡∏≠‡∏á‡∏ß‡πà‡∏≤‡∏á)"""
    a = t(a); b = t(b)
    if not a: return b
    if not b: return a
    return f"{a}  {b}"

# ===================== body paragraph =====================

def _reset_center_paragraph(p):
    """(‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏Å‡∏•‡∏≤‡∏á) ‡∏à‡∏±‡∏î‡∏Å‡∏∂‡πà‡∏á‡∏Å‡∏•‡∏≤‡∏á‡∏à‡∏£‡∏¥‡∏á‡πÅ‡∏•‡∏∞‡∏Å‡∏±‡∏ô indent ‡∏à‡∏≤‡∏Å Normal style"""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    return pf

# --- 1. ‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏ó‡∏µ‡πà‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç (‡∏£‡∏±‡∏ö chapter_no) ---

def add_picture_box_with_caption(
    doc: Document,
    abs_path: str,
    *,
    pic_name: str = "",
    chapter_no: int,
    run_no: int,
    width_cm: float = 5.8,
    height_cm: float = 4.8,
) -> None:
    """
    (‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏Å‡∏•‡∏≤‡∏á) ‡∏ß‡∏≤‡∏á‡∏£‡∏π‡∏õ ‡∏à‡∏±‡∏î‡∏Å‡∏∂‡πà‡∏á‡∏Å‡∏•‡∏≤‡∏á ‡πÅ‡∏•‡∏∞‡πÉ‡∏™‡πà‡πÅ‡∏Ñ‡∏õ‡∏ä‡∏±‡∏ô (‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏Å‡∏£‡∏≠‡∏ö)
    """
    # ‡∏¢‡πà‡∏≠‡∏´‡∏ô‡πâ‡∏≤‡∏£‡∏π‡∏õ
    p_img = doc.add_paragraph()
    _reset_center_paragraph(p_img)

    run = p_img.add_run()
    run.add_picture(abs_path, width=Cm(width_cm), height=Cm(height_cm))

    # ‡∏£‡∏∞‡∏¢‡∏∞‡∏Å‡πà‡∏≠‡∏ô/‡∏´‡∏•‡∏±‡∏á‡∏£‡∏π‡∏õ
    pf = p_img.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)

    # ‡πÅ‡∏Ñ‡∏õ‡∏ä‡∏±‡∏ô
    cap = doc.add_paragraph()
    _reset_center_paragraph(cap)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)

    # üîπ ‡∏ó‡∏≥‡∏ï‡∏±‡∏ß‡∏´‡∏ô‡∏≤‡πÄ‡∏â‡∏û‡∏≤‡∏∞ "‡∏†‡∏≤‡∏û‡∏ó‡∏µ‡πà x-y"
    prefix = f"‡∏†‡∏≤‡∏û‡∏ó‡∏µ‡πà {chapter_no}-{run_no}"
    r1 = cap.add_run(prefix + "  ")   # ‡∏°‡∏µ 2 ‡∏ä‡πà‡∏≠‡∏á‡∏ß‡πà‡∏≤‡∏á‡∏Ñ‡∏±‡πà‡∏ô
    r1.bold = True

    # üîπ ‡∏ä‡∏∑‡πà‡∏≠‡∏£‡∏π‡∏õ (pic_name) ‡πÄ‡∏õ‡πá‡∏ô‡∏ï‡∏±‡∏ß‡∏ö‡∏≤‡∏á
    if pic_name:
        r2 = cap.add_run(pic_name)
        r2.bold = False



# =============== ‡πÄ‡∏î‡∏¥‡∏ô‡∏ï‡πâ‡∏ô‡πÑ‡∏°‡πâ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏¢‡πà‡∏≠‡∏¢ ===============
def walk_item_tree(
    doc: Document, section_no: str, nodes: List[Dict[str, Any]], *,
    chapter_no: int, media_root: str, pic_counter: List[int], seen_pics: set[str],
    heading_func: callable, body_func: callable, caption_func: callable | None = None,
):
    if not isinstance(nodes, list): return
    caption_func = caption_func or body_func

    for i, node in enumerate(nodes):
        current_no = f"{section_no}.{i+1}" if section_no else f"{i+1}"
        title = t((node or {}).get("text"))
        paras = [t(x) for x in as_list((node or {}).get("paragraphs")) if t(x)]

        if current_no or title:
            heading_func(doc, current_no, title)

        for s in paras:
            body_func(doc, s)

        for pinfo in as_list((node or {}).get("pictures")):
            abs_path = resolve_image_path(pinfo, media_root)
            if not abs_path or abs_path in seen_pics: continue
            seen_pics.add(abs_path); pic_counter[0] += 1

            add_picture_box_with_caption(
                doc, abs_path, pic_name=t(pinfo.get("pic_name")),
                chapter_no=chapter_no, run_no=pic_counter[0],
            )
            for cap in as_list(pinfo.get("captions")):
                s = t(cap)
                if s: caption_func(doc, s)

        walk_item_tree(
            doc, current_no, as_list((node or {}).get("children")),
            chapter_no=chapter_no, media_root=media_root,
            pic_counter=pic_counter, seen_pics=seen_pics,
            heading_func=heading_func, body_func=body_func, caption_func=caption_func,
        )




def add_body_paragraph_style_1(doc: Document, text: str,disth: bool = False) -> None:
    """
    (‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏Å‡∏•‡∏≤‡∏á) ‡∏™‡πÑ‡∏ï‡∏•‡πå‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤: ‡∏¢‡πà‡∏≠‡∏´‡∏ô‡πâ‡∏≤ 1.85 ‡∏ã‡∏°.
    """
    p = add_wrapped_paragraph(doc, text, n=99999, disth=True, custom_tap=1.85)
    # (‡∏≠‡∏≤‡∏à‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡πÄ‡∏û‡∏¥‡πà‡∏°‡πÄ‡∏ï‡∏¥‡∏° ‡πÄ‡∏ä‡πà‡∏ô p.alignment = ... ‡∏ñ‡πâ‡∏≤ add_wrapped_paragraph ‡πÑ‡∏°‡πà‡πÑ‡∏î‡πâ‡∏ó‡∏≥)



def add_section_heading_level1_style_1(doc: Document, title_no: str, title: str) -> None:
    """
    (‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏Å‡∏•‡∏≤‡∏á) ‡∏™‡πÑ‡∏ï‡∏•‡πå‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏£‡∏∞‡∏î‡∏±‡∏ö 1: ‡∏´‡∏ô‡∏≤, ‡πÑ‡∏°‡πà‡∏¢‡πà‡∏≠‡∏´‡∏ô‡πâ‡∏≤, ‡∏´‡πà‡∏≤‡∏á 6pt
    """
    text = two_spaces_join(t(title_no), t(title)) # (‡∏ï‡πâ‡∏≠‡∏á import two_spaces_join)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = True
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)

def add_section_heading_level2_plus_style_1(doc: Document, title_no: str, title: str) -> None:
    """
    (‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏Å‡∏•‡∏≤‡∏á) ‡∏™‡πÑ‡∏ï‡∏•‡πå‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏£‡∏∞‡∏î‡∏±‡∏ö 2+: ‡πÑ‡∏°‡πà‡∏´‡∏ô‡∏≤ ‡∏¢‡πà‡∏≠ 0.75, ‡∏´‡πà‡∏≤‡∏á 3pt
    """
    text = two_spaces_join(t(title_no), t(title)) # (‡∏ï‡πâ‡∏≠‡∏á import two_spaces_join)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = False
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.75)
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)


# =================== NEW: mapped heading / thai alpha level-5 ===================

THAI_ALPHA = [
    "‡∏Å","‡∏Ç","‡∏Ñ","‡∏á","‡∏à","‡∏â","‡∏ä","‡∏ã","‡∏å","‡∏ç",
    "‡∏é","‡∏è","‡∏ê","‡∏ë","‡∏í","‡∏ì","‡∏î","‡∏ï","‡∏ñ","‡∏ó",
    "‡∏ò","‡∏ô","‡∏ö","‡∏õ","‡∏ú","‡∏ù","‡∏û","‡∏ü","‡∏†","‡∏°",
    "‡∏¢","‡∏£","‡∏•","‡∏ß","‡∏®","‡∏©","‡∏™","‡∏´","‡∏¨","‡∏≠","‡∏Æ"
]

def level_inside_chapter_py(number: str) -> int:
    """
    ‡πÅ‡∏õ‡∏•‡∏á '2.1'‚Üí1, '2.1.1'‚Üí2, '2.1.1.1'‚Üí3, '2.1.1.1.1'‚Üí4
    ‡∏Ñ‡∏∑‡∏≠ '‡∏ô‡∏±‡∏ö‡∏£‡∏∞‡∏î‡∏±‡∏ö‡∏´‡∏•‡∏±‡∏á‡πÄ‡∏•‡∏Ç‡∏ö‡∏ó' ‡πÉ‡∏´‡πâ‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥‡πÄ‡∏õ‡πá‡∏ô 1
    """
    s = str(number or "").strip()
    if not s:
        return 1
    parts = s.split(".")
    if len(parts) <= 1:
        return 1
    return max(1, len(parts) - 1)

def last_index_from_number_py(number: str) -> int:
    """
    ‡∏Ñ‡∏∑‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç‡∏ó‡πâ‡∏≤‡∏¢‡∏™‡∏∏‡∏î‡∏Ç‡∏≠‡∏á‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠ ‡πÄ‡∏ä‡πà‡∏ô '2.1.1.3' -> 3
    ‡∏´‡∏≤‡∏Å‡∏´‡∏≤‡πÑ‡∏°‡πà‡πÑ‡∏î‡πâ ‡∏Ñ‡∏∑‡∏ô 1
    """
    try:
        parts = str(number).split(".")
        tail = parts[-1] if parts else "1"
        return int(tail) if tail.isdigit() else 1
    except Exception:
        return 1

def add_section_heading_tap(doc, section_no: str = "", title: str = "",
                            left_cm: float = 0.0, bold: bool = False,
                            space_before_pt: int = 3, space_after_pt: int = 0):
    """
    ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏î‡πâ‡∏ß‡∏¢‡∏Å‡∏≤‡∏£‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏£‡∏∞‡∏¢‡∏∞‡πÄ‡∏¢‡∏∑‡πâ‡∏≠‡∏á‡∏ã‡πâ‡∏≤‡∏¢ (‡∏ã‡∏°.) ‡πÅ‡∏•‡∏∞‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏ï‡∏±‡∏ß‡∏´‡∏ô‡∏≤/‡∏ö‡∏≤‡∏á
    ‡πÑ‡∏°‡πà‡∏¢‡∏∏‡πà‡∏á‡∏Å‡∏±‡∏ö‡∏ü‡∏≠‡∏ô‡∏ï‡πå/‡∏Ç‡∏ô‡∏≤‡∏î ‡πÄ‡∏û‡∏£‡∏≤‡∏∞‡∏™‡∏∑‡∏ö‡∏ó‡∏≠‡∏î‡∏à‡∏≤‡∏Å doc_setup() ‡πÄ‡∏î‡∏¥‡∏°
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(left_cm)
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.keep_with_next = True

    text = two_spaces_join(section_no, title)
    r = p.add_run(text)
    r.font.bold = bool(bold)
    return p

def make_heading_tap_func_map(
    *,
    level_tap_cm: dict[int, float] | None = None,
    level_bold: dict[int, bool] | None = None,
    alpha_level: int = 5,
    alpha_list: list[str] | None = None,
    fallback_left_cm: float = 0.0,
):
    """
    ‡∏Ñ‡∏∑‡∏ô‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô heading_func(doc, number, title, level=?)
    - ‡∏£‡∏∞‡∏¢‡∏∞‡πÄ‡∏¢‡∏∑‡πâ‡∏≠‡∏á: ‡πÉ‡∏ä‡πâ level_tap_cm ‡∏ï‡∏≤‡∏°‡πÄ‡∏•‡πÄ‡∏ß‡∏• (‡πÄ‡∏ä‡πà‡∏ô {1:0, 2:0.75, 3:1.0, 4:2.0, 5:2.0})
    - ‡∏ï‡∏±‡∏ß‡∏´‡∏ô‡∏≤/‡∏ö‡∏≤‡∏á: ‡πÉ‡∏ä‡πâ level_bold ‡∏ï‡∏≤‡∏°‡πÄ‡∏•‡πÄ‡∏ß‡∏• (‡πÄ‡∏ä‡πà‡∏ô {1:True,2:True,3:False,4:False,5:False})
    - ‡∏ñ‡πâ‡∏≤ level == alpha_level (‡πÄ‡∏ä‡πà‡∏ô 5) ‡πÉ‡∏´‡πâ‡πÅ‡∏™‡∏î‡∏á prefix ‡πÄ‡∏õ‡πá‡∏ô‡∏≠‡∏±‡∏Å‡∏©‡∏£‡πÑ‡∏ó‡∏¢ '‡∏Å)' ‡πÅ‡∏ó‡∏ô‡πÄ‡∏•‡∏Ç‡∏ó‡πâ‡∏≤‡∏¢‡∏à‡∏£‡∏¥‡∏á
      ‡πÇ‡∏î‡∏¢ index ‡∏≠‡∏±‡∏Å‡∏©‡∏£‡∏≠‡πâ‡∏≤‡∏á‡∏≠‡∏¥‡∏á‡∏à‡∏≤‡∏Å‡πÄ‡∏•‡∏Ç‡∏ó‡πâ‡∏≤‡∏¢‡∏™‡∏∏‡∏î‡∏Ç‡∏≠‡∏á‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠
    """
    level_tap_cm = level_tap_cm or {}
    level_bold = level_bold or {}
    alpha_list = alpha_list or THAI_ALPHA

    def heading_func(d, number, text, level: int | None = None):
        lv = level_inside_chapter_py(number) if level is None else int(level) or level_inside_chapter_py(number)
        left_cm = level_tap_cm.get(lv, fallback_left_cm)
        is_bold = level_bold.get(lv, False)

        show_no = str(number)
        if alpha_level and lv == alpha_level:
            idx = max(1, last_index_from_number_py(number))
            letter = alpha_list[(idx - 1) % len(alpha_list)]
            show_no = f"{letter})"

        return add_section_heading_tap(d, show_no, text, left_cm=left_cm, bold=is_bold)

    return heading_func

# ----------function for views helpers ----------
def _t(v: Any) -> str:
    return (v or "").strip() if isinstance(v, str) else ""

def _safe_parse_list(raw: Any, fallback: list = None) -> list:
    if isinstance(raw, list):
        return raw
    try:
        data = json.loads(raw or "[]")
        return data if isinstance(data, list) else (fallback or [])
    except Exception:
        return fallback or []

def _default_sections_for_ui() -> List[Dict[str, Any]]:
    # UI ‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏£‡∏Å‡πÅ‡∏ö‡∏ö paragraphs, ‡∏ó‡∏µ‡πà‡πÄ‡∏´‡∏•‡∏∑‡∏≠‡πÉ‡∏ä‡πâ body+points
    out: List[Dict[str, Any]] = []
    for i, t in enumerate(DEFAULT_TITLES):
        if i == 0:
            out.append({"title": t, "paragraphs": [], "points": []})
        else:
            out.append({"title": t, "body": "", "points": []})
    return out

# ---------- 1) iterator ‡∏Å‡∏•‡∏≤‡∏á: ‡∏Ñ‡∏∑‡∏ô (title, body, mains) ----------
def iter_sections(sections_any: Any,*,first_section_mode: str = "paragraphs",  ) -> Iterable[Tuple[str, str, List[Dict[str, Any]]]]:
    """
    ‡∏£‡∏±‡∏ö schema ‡πÑ‡∏î‡πâ‡∏´‡∏•‡∏≤‡∏¢‡πÅ‡∏ö‡∏ö‡∏à‡∏≤‡∏Å UI/DB ‡πÅ‡∏•‡πâ‡∏ß 'normalize' ‡πÉ‡∏´‡πâ‡πÄ‡∏õ‡πá‡∏ô:
      - title: str
      - body: str
      - mains: List[{"text": str, "subs": List[str]}]
    ‡πÇ‡∏î‡∏¢:
      - ‡∏ñ‡πâ‡∏≤ first_section_mode == "paragraphs": ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏£‡∏Å‡πÉ‡∏ä‡πâ paragraphs (‡πÅ‡∏ï‡πà‡∏à‡∏∞‡∏Ñ‡∏∑‡∏ô body="" ‡πÅ‡∏•‡∏∞ mains=[])
      - ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏≠‡∏∑‡πà‡∏ô ‡πÜ ‡πÅ‡∏õ‡∏•‡∏á points->mains ‡πÑ‡∏î‡πâ
    """
    sections = _safe_parse_list(sections_any, [])
    for i, sec in enumerate(sections):
        if not isinstance(sec, dict):
            continue

        title = _t(sec.get("title") or sec.get("header") or sec.get("name"))
        body  = _t(sec.get("body") or sec.get("content") or sec.get("desc"))

        # mains ‡∏£‡∏±‡∏ö‡πÑ‡∏î‡πâ‡∏ó‡∏±‡πâ‡∏á mains/points/items
        points = sec.get("mains")
        if not isinstance(points, list):
            points = sec.get("points")
        if not isinstance(points, list):
            points = sec.get("items")
        points = points if isinstance(points, list) else []

        mains: List[Dict[str, Any]] = []
        for p in points:
            if isinstance(p, dict):
                text = _t(p.get("text") or p.get("main") or p.get("title") or p.get("name"))
                subs_in = p.get("subs") if isinstance(p.get("subs"), list) else []
                subs = [_t(s if isinstance(s, str) else (s.get("text") if isinstance(s, dict) else "")) for s in subs_in]
                mains.append({"text": text, "subs": [s for s in subs if s]})
            elif isinstance(p, str):
                mains.append({"text": _t(p), "subs": []})

        # ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏£‡∏Å‡πÅ‡∏ö‡∏ö paragraphs: ‡πÑ‡∏°‡πà‡∏Ñ‡∏∑‡∏ô mains; body ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏£‡∏Å‡∏Ñ‡∏á‡∏õ‡∏•‡πà‡∏≠‡∏¢‡∏ß‡πà‡∏≤‡∏á‡πÑ‡∏õ
        if i == 0 and first_section_mode == "paragraphs":
            yield (title, "", [])   # ‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤‡∏¢‡πà‡∏≠‡∏´‡∏ô‡πâ‡∏≤‡πÉ‡∏´‡∏ç‡πà‡πÉ‡∏´‡πâ‡πÑ‡∏õ‡∏≠‡πà‡∏≤‡∏ô‡∏à‡∏≤‡∏Å sec["paragraphs"] ‡∏ù‡∏±‡πà‡∏á‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£
        else:
            yield (title, body, mains)

# ---------- 2) UI -> DB (‡πÄ‡∏ã‡∏ü‡πÄ‡∏õ‡πá‡∏ô mains ‡∏Ñ‡∏á‡∏ó‡∏µ‡πà ‡πÉ‡∏ä‡πâ‡πÑ‡∏î‡πâ‡∏ó‡∏∏‡∏Å‡∏ö‡∏ó) ----------
def sections_db_from_ui(
    ui_any: Any,
    *,
    default_titles: Optional[List[str]] = None,
    first_section_mode: str = "paragraphs",  # "paragraphs" | "body"
) -> List[Dict[str, Any]]:
    ui = _safe_parse_list(ui_any, [])
    out: List[Dict[str, Any]] = []

    for i, sec in enumerate(ui):
        if not isinstance(sec, dict):
            continue
        title = _t(sec.get("title") or (default_titles[i] if default_titles and i < len(default_titles) else ""))

        if i == 0 and first_section_mode == "paragraphs":
            # ‡∏ö‡∏ó‡∏ó‡∏µ‡πà 1 (‡∏´‡∏£‡∏∑‡∏≠‡∏ö‡∏ó‡πÑ‡∏´‡∏ô‡∏Å‡πá‡∏ï‡∏≤‡∏°‡∏ó‡∏µ‡πà‡∏≠‡∏¢‡∏≤‡∏Å‡πÉ‡∏ä‡πâ paragraphs)
            paras_in = sec.get("paragraphs")
            paras = []
            if isinstance(paras_in, list):
                for it in paras_in:
                    s = _t(it if isinstance(it, str) else (it.get("text") if isinstance(it, dict) else ""))
                    if s: paras.append(s)
            body = "\n\n".join(paras)  # ‡πÄ‡∏Å‡πá‡∏ö body ‡πÑ‡∏ß‡πâ‡πÄ‡∏ú‡∏∑‡πà‡∏≠‡∏£‡∏∞‡∏ö‡∏ö‡πÄ‡∏Å‡πà‡∏≤
            out.append({"title": title, "paragraphs": paras, "body": body, "mains": []})
            continue

        # ‡∏ó‡∏µ‡πà‡πÄ‡∏´‡∏•‡∏∑‡∏≠: normalize ‡πÄ‡∏õ‡πá‡∏ô mains ‡πÄ‡∏™‡∏°‡∏≠ (‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö points/mains)
        body = _t(sec.get("body"))
        mains_out: List[Dict[str, Any]] = []

        if isinstance(sec.get("mains"), list) and sec["mains"]:
            for m in sec["mains"]:
                if not isinstance(m, dict): continue
                text = _t(m.get("text") or m.get("title") or m.get("name"))
                subs_in = m.get("subs") if isinstance(m.get("subs"), list) else []
                subs = [_t(s if isinstance(s, str) else (s.get("text") if isinstance(s, dict) else "")) for s in subs_in]
                mains_out.append({"text": text, "subs": [x for x in subs if x]})
        elif isinstance(sec.get("points"), list):
            for p in sec["points"]:
                if isinstance(p, dict):
                    text = _t(p.get("main") or p.get("text") or p.get("title"))
                    subs_in = p.get("subs") if isinstance(p.get("subs"), list) else []
                    subs = [_t(s if isinstance(s, str) else (s.get("text") if isinstance(s, dict) else "")) for s in subs_in]
                    mains_out.append({"text": text, "subs": [x for x in subs if x]})
                elif isinstance(p, str):
                    txt = _t(p)
                    if txt: mains_out.append({"text": txt, "subs": []})

        out.append({"title": title, "body": body, "mains": mains_out})

    return out

# ---------- 3) DB -> UI (‡∏Ñ‡∏∑‡∏ô schema ‡πÄ‡∏î‡∏¥‡∏°‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ü‡∏≠‡∏£‡πå‡∏°: ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏£‡∏Å paragraphs, ‡∏ó‡∏µ‡πà‡πÄ‡∏´‡∏•‡∏∑‡∏≠ body+points) ----------
def sections_ui_from_db(
    db_any: Any,
    *,
    default_titles: Optional[List[str]] = None,
    first_section_mode: str = "paragraphs",
) -> List[Dict[str, Any]]:
    db = _safe_parse_list(db_any, [])

    # ‡πÄ‡∏ï‡∏£‡∏µ‡∏¢‡∏°‡πÇ‡∏Ñ‡∏£‡∏á UI ‡∏ï‡∏≤‡∏°‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠
    ui: List[Dict[str, Any]] = []
    total = max(len(db), len(default_titles or []))
    for i in range(total or 1):
        t = (default_titles[i] if default_titles and i < len(default_titles) else "")
        if i == 0 and first_section_mode == "paragraphs":
            ui.append({"title": t, "paragraphs": [], "points": []})
        else:
            ui.append({"title": t, "body": "", "points": []})

    for i, sec in enumerate(db):
        if not isinstance(sec, dict) or i >= len(ui):
            continue
        title = _t(sec.get("title") or (default_titles[i] if default_titles and i < len(default_titles) else ""))

        if i == 0 and first_section_mode == "paragraphs":
            paras = []
            if isinstance(sec.get("paragraphs"), list) and sec["paragraphs"]:
                for it in sec["paragraphs"]:
                    s = _t(it if isinstance(it, str) else (it.get("text") if isinstance(it, dict) else ""))
                    if s: paras.append(s)
            else:
                body = _t(sec.get("body"))
                if body:
                    paras = [p.strip() for p in body.replace("\r\n", "\n").split("\n\n") if p.strip()]
            ui[i]["title"] = title
            ui[i]["paragraphs"] = paras
            continue

        ui[i]["title"] = title
        ui[i]["body"] = _t(sec.get("body"))
        points = []
        mains = sec.get("mains") if isinstance(sec.get("mains"), list) else []
        for m in mains:
            if not isinstance(m, dict): 
                if isinstance(m, str):
                    points.append({"main": _t(m), "subs": []})
                continue
            main_text = _t(m.get("text") or m.get("title") or m.get("name"))
            subs_out = []
            subs = m.get("subs") if isinstance(m.get("subs"), list) else []
            for s in subs:
                subs_out.append(_t(s if isinstance(s, str) else (s.get("text") if isinstance(s, dict) else "")))
            points.append({"main": main_text, "subs": [x for x in subs_out if x]})
        ui[i]["points"] = points

    return ui

# ---------- 4) ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÄ‡∏≠‡∏ô‡∏à‡∏¥‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£: ‡∏Ñ‡πâ‡∏≥‡∏õ‡∏£‡∏∞‡∏Å‡∏±‡∏ô‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏Å‡πà‡∏≠‡∏ô render ----------
def sections_doc_safe(
    sections_any: Any,
    *,
    default_titles: List[str],
    first_section_mode: str = "paragraphs",
) -> List[Dict[str, Any]]:
    
    raw = _safe_parse_list(sections_any, [])
    want_n = len(default_titles or [])

    out: List[Dict[str, Any]] = []
    for i in range(want_n or 1):
        t = default_titles[i] if default_titles and i < len(default_titles) else ""
        if i == 0 and first_section_mode == "paragraphs":
            out.append({"title": t, "body": "", "paragraphs": [], "mains": []})
        else:
            out.append({"title": t, "body": "", "mains": []})

    for i in range(min(len(raw), len(out))):
        src = raw[i] if isinstance(raw[i], dict) else {}
        title = _t(src.get("title") or out[i]["title"])
        body  = _t(src.get("body"))

        if i == 0 and first_section_mode == "paragraphs":
            paras = []
            if isinstance(src.get("paragraphs"), list):
                for p in src["paragraphs"]:
                    s = _t(p if isinstance(p, str) else (p.get("text") if isinstance(p, dict) else ""))
                    if s: paras.append(s)
            if not paras and body:
                paras = [p.strip() for p in body.replace("\r\n","\n").split("\n\n") if p.strip()]
            out[i].update({"title": title, "body": body, "paragraphs": paras, "mains": []})
        else:
            mains_out = []
            mains_in = src.get("mains") if isinstance(src.get("mains"), list) else []
            for m in mains_in:
                if isinstance(m, dict):
                    text = _t(m.get("text") or m.get("title") or m.get("name"))
                    subs_in = m.get("subs") if isinstance(m.get("subs"), list) else []
                    subs = [_t(s if isinstance(s, str) else (s.get("text") if isinstance(s, dict) else "")) for s in subs_in]
                    mains_out.append({"text": text, "subs": [x for x in subs if x]})
                elif isinstance(m, str):
                    mains_out.append({"text": _t(m), "subs": []})
            out[i].update({"title": title, "body": body, "mains": mains_out})

    return out
# wrapped_paragraph ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Ñ‡∏≥‡∏≠‡∏ò‡∏¥‡∏ö‡∏≤‡∏¢‡πÉ‡∏ï‡πâ‡∏£‡∏π‡∏õ
def add_intro_caption_paragraph(doc: Document, text: str) -> None:
    add_wrapped_paragraph(doc, text, n=99999, custom_tap=0.75,disth=True)